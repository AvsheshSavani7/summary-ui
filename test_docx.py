from docx import Document
from docx.shared import Pt, Inches, RGBColor
from collections import defaultdict
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SECTION_ORDER = ["Conditions", "Termination", "Consideration", "Other"]


def add_tab_stop(paragraph, position_inches):
    pPr = paragraph._element.get_or_add_pPr()
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        tabs = OxmlElement('w:tabs')
        pPr.append(tabs)

    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'left')
    tab.set(qn('w:pos'), str(int(position_inches * 1440)))  # inches to twips
    tabs.append(tab)


def write_docx_summary(summaries, output_path="test_output.docx"):
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Aptos"
    font.size = Pt(10.5)

    # Set margins to 1"
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Merger Agreement Clause Summaries")
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    title_run.font.name = "Aptos"
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph("")

    # Group summaries
    grouped = defaultdict(lambda: defaultdict(list))
    for s in summaries:
        s_type = s.get("summary_type", "fulsome").capitalize()
        s_section = s.get("summary_display_section", "General")
        grouped[s_type][s_section].append(s)

    for i, s_type in enumerate(["Concise", "Fulsome"]):
        if s_type in grouped:
            if i == 1:
                doc.add_paragraph("")

            p = doc.add_paragraph()
            run = p.add_run(f"{s_type} Summary")
            run.underline = True
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Aptos"
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.space_after = Pt(4)

            for section in SECTION_ORDER + [k for k in grouped[s_type] if k not in SECTION_ORDER]:
                if section in grouped[s_type]:
                    for s in sorted(grouped[s_type][section], key=lambda x: x.get("summary_rank", 999)):
                        # Main bullet line (+)
                        bullet_para = doc.add_paragraph()
                        bullet_para.paragraph_format.left_indent = Inches(0.25)
                        bullet_para.paragraph_format.first_line_indent = - \
                            Inches(0.25)
                        bullet_para.paragraph_format.space_after = Pt(4)
                        bullet_para.paragraph_format.line_spacing = 1.16
                        add_tab_stop(bullet_para, 0.25)

                        bullet_run = bullet_para.add_run("+\t")
                        bullet_run.font.name = "Aptos"
                        bullet_run.font.size = Pt(10.5)
                        bullet_run.font.color.rgb = RGBColor(0, 0, 0)

                        text_run = bullet_para.add_run(s["output"])
                        text_run.font.name = "Aptos"
                        text_run.font.size = Pt(10.5)
                        text_run.font.color.rgb = RGBColor(0, 0, 0)

                        # Reference line with hollow circle bullet
                        if s.get("references"):
                            ref_para = doc.add_paragraph()
                            ref_para.paragraph_format.left_indent = Inches(1.0)
                            ref_para.paragraph_format.first_line_indent = - \
                                Inches(0.25)
                            ref_para.paragraph_format.space_after = Pt(6)
                            ref_para.paragraph_format.line_spacing = 1.16
                            add_tab_stop(ref_para, 1.0)

                            ref_bullet = ref_para.add_run(
                                "○\t")  # True hollow circle
                            ref_bullet.font.name = "Aptos"
                            # Matches 10pt text height
                            ref_bullet.font.size = Pt(8)
                            ref_bullet.font.color.rgb = RGBColor(0, 0, 0)

                            ref_text = ref_para.add_run(
                                "References: " + "; ".join(s["references"]))
                            ref_text.font.name = "Aptos"
                            ref_text.font.size = Pt(10)
                            ref_text.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(output_path)
    print(f"\n✅ DOCX summary written to: {output_path}")


# Sample content
if __name__ == "__main__":
    sample = [
        {
            "summary_type": "concise",
            "summary_display_section": "Conditions",
            "output": "Each party must use reasonable best efforts to obtain regulatory approvals, including any second requests, competition filings, and government notifications, without undue delay or allocation shifts.",
            "references": ["7.01(b)"],
            "summary_rank": 1
        },
        {
            "summary_type": "fulsome",
            "summary_display_section": "Conditions",
            "output": "The parties shall use their respective reasonable best efforts to effect the transactions contemplated, including making all required antitrust filings, cooperating with each other in responding to regulatory inquiries, and pursuing litigation or appeals if necessary, unless doing so would materially and adversely affect the business or valuation of either party or its subsidiaries. In no event shall either party be required to divest any material portion of its business unless the aggregate value of such divestitures falls below the divestiture cap of $250 million, as mutually agreed.",
            "references": ["7.02(a)", "7.03(c)"],
            "summary_rank": 2
        }
    ]

    write_docx_summary(sample)
