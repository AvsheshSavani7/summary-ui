import openai
import os
import re
import sys
import json
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from collections import defaultdict
import datetime
import dateutil.parser

# summary_engine.py
RUN_CONCISE_SUMMARIES = True
RUN_FULSOME_SUMMARIES = True

# =========================
# LLM Setup
# =========================


def load_api_key():
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY not found in .env file")
    return api_key


openai.api_key = load_api_key()


def get_summary_mode_toggles():
    return RUN_CONCISE_SUMMARIES, RUN_FULSOME_SUMMARIES


def call_llm(prompt_text, model="gpt-4", temperature=0):
    response = openai.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are a legal summarization assistant."},
            {"role": "user", "content": prompt_text}
        ],
        temperature=temperature
    )
    return response.choices[0].message.content.strip()

# =========================
# Utility to Traverse Nested Data
# =========================


def get_nested_value(data, path):
    keys = path.split(".")
    for key in keys:
        if not isinstance(data, dict) or key not in data:
            return None
        data = data[key]
    return data

# =========================
# Utility to Extract Short Reference
# =========================


def shorten_reference(ref_string):
    ref_string = ref_string.replace("ARTICLE Article", "Article")
    match = re.search(
        r"(Section\\s\\d+(\\.\\d+)?([a-z])?)", ref_string, re.IGNORECASE)
    if match:
        return match.group(1).strip()
    return ref_string.strip()


def extract_short_reference(references, data, fallback_to_section=True):
    short_refs = []
    for ref_path in references:
        val = get_nested_value(data, ref_path)
        if val is None:
            continue
        if ref_path.endswith("short_reference"):
            short_refs.append(val)
        elif fallback_to_section and isinstance(val, str):
            short_refs.append(shorten_reference(val))
        else:
            short_refs.append(val)
    return short_refs

# =========================
# Evaluate One Condition or Group
# =========================


def evaluate_condition_branch(condition, data):
    if "conditions" in condition:
        run_if = condition.get("run_if", "all")
        subresults = [evaluate_condition_branch(
            sub, data) for sub in condition["conditions"]]
        should_run = all("triggered" in res for res in subresults) if run_if == "all" else any(
            "triggered" in res for res in subresults)

        if should_run:
            output = {"triggered": True}
            for res in subresults:
                if "add_to_prompt" in res:
                    output.setdefault("add_to_prompt", {}).update(
                        res["add_to_prompt"])
                if "add_references" in res:
                    output.setdefault("add_references", []).extend(
                        res["add_references"])
            return output
        return {}

    value = get_nested_value(data, condition["if"])
    result = None
    output = {}

    if condition["type"] == "boolean":
        result = bool(value)

    elif condition["type"] == "enum":
        if isinstance(value, list):
            value_normalized = ", ".join(str(v) for v in value).strip()
        else:
            value_normalized = (value or "").strip()
        enum_cases = condition.get("enum_cases", {})
        branch = enum_cases.get(value_normalized, condition.get("default", {}))
        if branch:
            output["triggered"] = True

        if "text_output" in branch:
            output["text_output"] = branch["text_output"]

        if "add_to_prompt" in branch:
            resolved_prompt_fields = resolve_prompt_fields(
                branch["add_to_prompt"], data)
            output.setdefault("add_to_prompt", {}).update(
                resolved_prompt_fields)

        if "add_references" in branch:
            output.setdefault("add_references", []).extend(
                branch["add_references"])

        return output

    elif condition["type"] == "number":
        comparator = condition.get("comparator", "==")
        compare_value = condition.get("compare_to", condition.get("value"))
        if value is None or compare_value is None:
            result = False
        elif comparator == "==":
            result = value == compare_value
        elif comparator == ">":
            result = value > compare_value
        elif comparator == "<":
            result = value < compare_value
        elif comparator == ">=":
            result = value >= compare_value
        elif comparator == "<=":
            result = value <= compare_value
        else:
            result = False

    elif condition["type"] == "non_empty":
        result = value is not None and value != ""

    if result is True and "true" in condition:
        branch = condition["true"]
    elif result is False and "false" in condition:
        branch = condition["false"]
    elif "default" in condition:
        branch = condition["default"]
    else:
        branch = {}

    if branch:
        output["triggered"] = True
    if "text_output" in branch:
        output["text_output"] = branch["text_output"]
    if "add_to_prompt" in branch:
        resolved_prompt_fields = resolve_prompt_fields(
            branch["add_to_prompt"], data)
        output.setdefault("add_to_prompt", {}).update(resolved_prompt_fields)
    if "add_references" in branch:
        output.setdefault("add_references", []).extend(
            branch["add_references"])

    return output


def resolve_prompt_fields(prompt_dict, data):
    resolved = {}
    for k, v in prompt_dict.items():
        if isinstance(v, str) and v.startswith("{{") and v.endswith("}}"):
            key_path = v[2:-2].strip()  # Strip the double braces
            val = get_nested_value(data, key_path)

            if val and isinstance(val, str):
                match = re.search(
                    r"(\d+)\s+Business Days after.*(date of the Agreement|Signing Date)", val, re.IGNORECASE)
                if match:
                    num_days = int(match.group(1))
                    base_str = get_nested_value(
                        data, "timeline.agreement_signing_date.agreement_signing_date.clause_text")
                    try:
                        base_date = dateutil.parser.parse(base_str)
                        computed = add_business_days(base_date, num_days)
                        resolved[k] = computed.strftime("%B %d, %Y")
                    except Exception:
                        resolved[k] = val
                else:
                    resolved[k] = val
            else:
                resolved[k] = val if val is not None else ""
        else:
            # Treat as literal value
            resolved[k] = v
    return resolved


def normalize_to_string_list(value):
    if isinstance(value, str):
        return [value]
    elif isinstance(value, dict):
        return [str(value)]
    elif isinstance(value, list):
        return [str(item) if not isinstance(item, str) else item for item in value]
    else:
        return [str(value)]  # fallback for other types

# =========================
# Process a Single Clause Config
# =========================


def process_clause_config(clause_config, schema_data):

    prompt_fields = {}
    references = []
    final_text_output = None
    for cond in clause_config.get("conditions", []):
        result = evaluate_condition_branch(cond, schema_data)
        if "add_to_prompt" in result:
            for k, v in result["add_to_prompt"].items():
                if k not in prompt_fields:
                    prompt_fields[k] = v
                else:
                    if isinstance(prompt_fields[k], list):
                        prompt_fields[k].append(v)
                    else:
                        prompt_fields[k] = [prompt_fields[k], v]
        if "add_references" in result:
            references.extend(result["add_references"])
        if "text_output" in result:
            final_text_output = result["text_output"]
    # Handle list-based prompt fields
    for k, v in prompt_fields.items():
        if isinstance(v, list):
            join_type = clause_config.get("join_type", "bullets")
            if join_type == "bullets":
                print(f"Processing field: {k}")
                print(f"Value (v): {v}")
                print(
                    f"Type of first item: {type(v[0]) if isinstance(v, list) and v else 'N/A'}")
                prompt_fields[k] = "\n- " + \
                    "\n- ".join(normalize_to_string_list(v))
            elif join_type == "sentences":
                prompt_fields[k] = " ".join(normalize_to_string_list(v))
            else:
                prompt_fields[k] = "\n".join(normalize_to_string_list(v))
    references.extend(clause_config.get("reference_fields", []))
    references = list(set(references))
    if clause_config.get("use_short_reference", True):
        short_refs = extract_short_reference(
            references, schema_data, fallback_to_section=True)
    else:
        short_refs = [
            get_nested_value(schema_data, path)
            for path in references
            if get_nested_value(schema_data, path)
        ]
    # If prompt can be built
    if prompt_fields and "prompt_template" in clause_config:
        try:
            prompt = clause_config["prompt_template"].format(**prompt_fields)
            if clause_config.get("max_words"):
                prompt += f"\n\nLimit the response to {clause_config['max_words']} words."
            if clause_config.get("format_style"):
                prompt += f"\n\nFormat the response in a {clause_config['format_style']} style."
        except KeyError as e:
            if "fallback_prompt" in clause_config:
                prompt = clause_config["fallback_prompt"]
            else:
                prompt = f"[Missing field {str(e)} for prompt generation]"
        llm_result = call_llm(prompt)
        return {
            "output": llm_result,
            "references": short_refs,
            "used_prompt": prompt if clause_config.get("view_prompt", False) else None,
            "summary_type": clause_config.get("summary_type"),
            "format_style": clause_config.get("format_style"),
            "summary_display_section": clause_config.get("summary_display_section"),
            # "summary_display_sub_section" : clause_config.get("summary_display_sub_section"),
            "summary_rank": clause_config.get("summary_rank"),
            "max_words": clause_config.get("max_words")
        }
    # If no prompt was built, use fallback text_output
    if final_text_output:
        return {
            "output": final_text_output,
            "references": short_refs,
            "used_prompt": None,
            "summary_type": clause_config.get("summary_type"),
            "format_style": clause_config.get("format_style"),
            "summary_display_section": clause_config.get("summary_display_section"),
            # "summary_display_sub_section" : clause_config.get("summary_display_sub_section"),
            "summary_rank": clause_config.get("summary_rank"),
            "max_words": clause_config.get("max_words")
        }
    return {
        "output": "No output generated.",
        "references": short_refs,
        "used_prompt": None,
        "summary_type": clause_config.get("summary_type"),
        "format_style": clause_config.get("format_style"),
        "summary_display_section": clause_config.get("summary_display_section"),
        # "summary_display_sub_section" : clause_config.get("summary_display_sub_section"),
        "summary_rank": clause_config.get("summary_rank"),
        "max_words": clause_config.get("max_words")
    }


# =========================
# DOCX Writer
# =========================
SECTION_ORDER = ["Deal Structure", "Termination",
                 "Conditions", "Non Solicitation", "Regulatory", "General"]


def add_tab_stop(paragraph, position_inches):
    pPr = paragraph._element.get_or_add_pPr()
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        tabs = OxmlElement('w:tabs')
        pPr.append(tabs)

    # Avoid adding duplicate tab stops
    position_twips = str(int(position_inches * 1440))
    for existing_tab in tabs.findall(qn('w:tab')):
        if existing_tab.get(qn('w:pos')) == position_twips:
            return  # Already exists

    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'left')
    tab.set(qn('w:pos'), position_twips)
    tabs.append(tab)


def write_docx_summary(summaries, output_path, RUN_CONCISE_SUMMARIES, RUN_FULSOME_SUMMARIES):

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Aptos"
    font.size = Pt(10.5)

    # Set margins to 1" all around
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Merger Agreement Clause Summaries")
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    title_run.font.name = "Aptos"
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph("")

    # Concise / Fulsome grouping and content writing
    enabled_summary_types = []
    if RUN_CONCISE_SUMMARIES:
        enabled_summary_types.append("Concise")
    if RUN_FULSOME_SUMMARIES:
        enabled_summary_types.append("Fulsome")

    for s_type in enabled_summary_types:
        # Header
        p = doc.add_paragraph()
        run = p.add_run(f"{s_type} Summary")
        run.underline = True
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Aptos"
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_after = Pt(4)

        already_print = []
        for s in [summary for summary in summaries if summary.get("summary_type") == s_type]:
            if s.get("summary_display_section") not in already_print:
                doc.add_heading(s.get("summary_display_section"), level=2)
                already_print.append(s.get("summary_display_section"))

            bullet_para = doc.add_paragraph()
            bullet_para.paragraph_format.left_indent = Inches(0.25)
            bullet_para.paragraph_format.first_line_indent = -Inches(0.25)
            bullet_para.paragraph_format.line_spacing = 1.16
            bullet_para.paragraph_format.space_before = Pt(1)
            bullet_para.paragraph_format.space_after = Pt(0)
            add_tab_stop(bullet_para, 0.25)

            bullet_run = bullet_para.add_run("+\t")
            bullet_run.font.name = "Aptos"
            bullet_run.font.size = Pt(10.5)

            clean_output = s["output"].strip().strip('"').strip("'")
            text_run = bullet_para.add_run(clean_output)
            text_run.font.name = "Aptos"
            text_run.font.size = Pt(10.5)

            if s.get("references") and s["references"][0] != '':
                ref_para = doc.add_paragraph()
                ref_para.paragraph_format.left_indent = Inches(1.0)
                ref_para.paragraph_format.first_line_indent = -Inches(0.25)
                ref_para.paragraph_format.line_spacing = 1
                add_tab_stop(ref_para, 1.0)

                ref_bullet = ref_para.add_run("○\t")
                ref_bullet.font.name = "Aptos"
                ref_bullet.font.size = Pt(8)
                ref_bullet.font.color.rgb = RGBColor(0, 0, 0)

                ref_text = ref_para.add_run(
                    "References: " + "; ".join(s["references"]))
                ref_text.font.name = "Aptos"
                ref_text.font.size = Pt(10)

    doc.save(output_path)
    print(f"\n✅ DOCX summary written to: {output_path}")


# =========================
# Main Test Block
# =========================
if __name__ == "__main__":
    sys.path.append(os.path.dirname(__file__))
    from clause_configs.ordinary_course_config import ORDINARY_COURSE_CLAUSES

    CLAUSE_CONFIG = ORDINARY_COURSE_CLAUSES

    with open("schema.json") as f:
        EXAMPLE_SCHEMA_DATA = json.load(f)

    print("Loaded clause configs:", list(CLAUSE_CONFIG.keys()))
    summary_outputs = []

    print(f"📝 Writing {len(summary_outputs)} summaries to DOCX")

    for clause_name, clause_config in CLAUSE_CONFIG.items():
        # Assert OFF clauses were filtered upstream (or flag if not)
        assert clause_config.get(
            "summary_type") != "OFF", f"OFF clause was not skipped: {clause_name}"

        print(f"→ Evaluating: {clause_name}")
        result = process_clause_config(clause_config, EXAMPLE_SCHEMA_DATA)
        print(f"→ Output preview: {result['output'][:100]}")
        if result["output"] and result["output"] != "No output generated.":
            filtered_result = result.copy()
            if not clause_config.get("view_prompt", False):
                filtered_result.pop("used_prompt", None)

            summary_outputs.append({
                "clause_name": clause_name,
                **filtered_result
            })

    write_docx_summary(summary_outputs)
