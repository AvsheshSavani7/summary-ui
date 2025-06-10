import streamlit as st
import os
import json
import importlib
import re
import hashlib
import sys
import boto3
from io import StringIO
import tempfile
from dotenv import load_dotenv
from summary_engine import process_clause_config, write_docx_summary
from docx import Document
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

# Configure AWS S3
s3_client = boto3.client(
    's3',
    aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
    aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY'),
    region_name=os.getenv('AWS_REGION')
)
S3_BUCKET = os.getenv('AWS_S3_BUCKET')

# Configure the page
st.set_page_config(
    page_title="Prompt Template Editor",
    page_icon="✏️",
    layout="wide"
)

# Initialize session state for storing edited templates and selected template
if 'edited_templates' not in st.session_state:
    st.session_state.edited_templates = {}

if 'selected_template' not in st.session_state:
    st.session_state.selected_template = None

if 'template_lines' not in st.session_state:
    st.session_state.template_lines = {}

if 'line_order' not in st.session_state:
    st.session_state.line_order = {}

if 'needs_rerun' not in st.session_state:
    st.session_state.needs_rerun = False

# Add custom CSS for drag and drop functionality
st.markdown("""
<style>
.draggable-line {
    cursor: move;
    padding: 5px;
    margin: 2px 0;
    background-color: #f0f2f6;
    border: 1px solid #e0e0e0;
    border-radius: 4px;
}
/* Reduce sidebar width */
[data-testid="stSidebar"] {
    min-width: 25vw !important;
    max-width: 25vw !important;
}
[data-testid="stSidebar"] > div:first-child {
    width: 25vw !important;
}
.draggable-line:hover {
    background-color: #e6e9ef;
}
.line-controls {
    display: flex;
    align-items: center;
    gap: 5px;
}
.line-number {
    color: #666;
    font-size: 12px;
    width: 30px;
}
/* Reduce main section padding */
.block-container {
    padding-top: 1rem !important;
    padding-bottom: 1rem !important;
    padding-left: 1rem !important;
    padding-right: 1rem !important;
}
/* Reduce padding in elements */
.stTextArea {
    padding-top: 0.2rem !important;
    padding-bottom: 0.2rem !important;
}
.st-emotion-cache-16idsys {
    padding-top: 0.2rem !important;
    padding-bottom: 0.2rem !important;
}
/* Reduce spacing between elements */
div.stButton > button {
    margin-top: 0.2rem !important;
    margin-bottom: 0.2rem !important;
}
/* Reduce markdown spacing */
.element-container {
    margin-top: 0.2rem !important;
    margin-bottom: 0.2rem !important;
}
/* Reduce expander padding */
.streamlit-expanderHeader {
    padding-top: 0.2rem !important;
    padding-bottom: 0.2rem !important;
}
.streamlit-expanderContent {
    padding-top: 0.2rem !important;
    padding-bottom: 0.2rem !important;
}
/* Reduce code block padding */
.stCodeBlock {
    padding: 0rem !important;
}
            
/* remove gap between elements */
.st-emotion-cache-q5aj5z {
   gap: 0.1rem !important;
}

/* Styles for editable group header */
.group-header {
    display: flex;
    align-items: center;
    justify-content: space-between;
    margin-bottom: 0.1rem;
}

.group-header-text {
    margin: 0 !important;
    padding: 0 !important;
}

.delete-group-btn {
    padding: 0.2rem 0.5rem !important;
    height: auto !important;
    min-height: unset !important;
    line-height: 1 !important;
}

/* Move column styling */
[data-testid="column"]:first-child {
    display: flex !important;
    flex-direction: column !important;
    justify-content: space-between !important;
    min-height: 68px !important;
}

/* Adjust button spacing in move column */
[data-testid="column"]:first-child .stButton {
    margin: 0 !important;
}

[data-testid="column"]:first-child button {
    padding: 0.2rem 0.5rem !important;
    min-height: unset !important;
    height: auto !important;
    line-height: 1 !important;
}
           
           
/* Hide sidebar header */
[data-testid="stSidebarHeader"] {
    display: none !important;
}

/* Style hr inside markdown container */
[data-testid="stMarkdownContainer"] hr {
   margin-top: 0.2rem !important;
   margin-bottom: 0.2rem !important;
}
/* Sidebar gap */
[data-testid="stVerticalBlock"] {
 gap: 0rem !important;
}


</style>
        
""", unsafe_allow_html=True)


def delete_line(template_name, index):
    """Delete line at specified index"""
    lines = st.session_state.template_lines[template_name]
    order = st.session_state.line_order[template_name]

    if len(lines) > 1:  # Prevent deleting the last line
        del lines[index]
        # Update order
        new_order = list(range(len(lines)))
        st.session_state.line_order[template_name] = new_order


def on_text_change(template_name, group_index):
    """Callback function for text area changes"""
    key = f"{template_name}_group_{group_index}"
    if key in st.session_state:
        edited_text = st.session_state[key]
        new_lines = [line for line in edited_text.split("\n") if line.strip()]

        if new_lines:  # If there are non-empty lines
            st.session_state[f"{template_name}_groups"][group_index] = new_lines

            # Update complete template
            all_lines = []
            for idx in range(len(st.session_state[f"{template_name}_groups"])):
                if idx not in st.session_state[f"{template_name}_deleted_groups"]:
                    all_lines.extend(
                        st.session_state[f"{template_name}_groups"][idx])
            st.session_state.edited_templates[template_name] = "\n".join(
                all_lines)
        else:  # If all lines were deleted
            st.session_state[f"{template_name}_groups"][group_index] = []
            st.rerun()


def render_editable_groups(template_name, groups, dynamic_values):
    """Render all groups with the ability to move them, while maintaining editability status"""

    # Initialize session state for all groups if not exists
    if f"{template_name}_groups" not in st.session_state:
        st.session_state[f"{template_name}_groups"] = []
        # Track editability
        st.session_state[f"{template_name}_groups_editable"] = []
        # Track deleted groups
        st.session_state[f"{template_name}_deleted_groups"] = set()
        for group in groups:
            st.session_state[f"{template_name}_groups"].append(group["lines"])
            st.session_state[f"{template_name}_groups_editable"].append(
                group["editable"])

    # Container for all groups
    groups_container = st.container()

    with groups_container:
        edited_lines = []
        groups_to_remove = set()  # Track indices of groups to remove

        # Process each group
        for group_index in range(len(st.session_state[f"{template_name}_groups"])):
            # Skip if group was previously deleted
            if group_index in st.session_state[f"{template_name}_deleted_groups"]:
                continue

            current_group = st.session_state[f"{template_name}_groups"][group_index]
            is_editable = st.session_state[f"{template_name}_groups_editable"][group_index]

            # Skip empty non-editable groups
            if not is_editable and not current_group:
                groups_to_remove.add(group_index)
                continue

            # Create columns for group controls and content
            move_col, content_col = st.columns([0.05, 0.95])

            with move_col:
                # Move up button
                if group_index > 0:
                    if st.button("↑", key=f"group_up_{group_index}", help="Move group up"):
                        # Find next valid index up
                        next_index = group_index - 1
                        while next_index in st.session_state[f"{template_name}_deleted_groups"] and next_index > 0:
                            next_index -= 1

                        if next_index >= 0 and next_index not in st.session_state[f"{template_name}_deleted_groups"]:
                            # Swap groups
                            st.session_state[f"{template_name}_groups"][group_index], \
                                st.session_state[f"{template_name}_groups"][next_index] = \
                                st.session_state[f"{template_name}_groups"][next_index], \
                                st.session_state[f"{template_name}_groups"][group_index]

                            # Swap editability flags
                            st.session_state[f"{template_name}_groups_editable"][group_index], \
                                st.session_state[f"{template_name}_groups_editable"][next_index] = \
                                st.session_state[f"{template_name}_groups_editable"][next_index], \
                                st.session_state[f"{template_name}_groups_editable"][group_index]
                            st.rerun()

                # Move down button
                if group_index < len(st.session_state[f"{template_name}_groups"]) - 1:
                    if st.button("↓", key=f"group_down_{group_index}", help="Move group down"):
                        # Find next valid index down
                        next_index = group_index + 1
                        while next_index in st.session_state[f"{template_name}_deleted_groups"] and next_index < len(st.session_state[f"{template_name}_groups"]) - 1:
                            next_index += 1

                        if next_index < len(st.session_state[f"{template_name}_groups"]) and next_index not in st.session_state[f"{template_name}_deleted_groups"]:
                            # Swap groups
                            st.session_state[f"{template_name}_groups"][group_index], \
                                st.session_state[f"{template_name}_groups"][next_index] = \
                                st.session_state[f"{template_name}_groups"][next_index], \
                                st.session_state[f"{template_name}_groups"][group_index]

                            # Swap editability flags
                            st.session_state[f"{template_name}_groups_editable"][group_index], \
                                st.session_state[f"{template_name}_groups_editable"][next_index] = \
                                st.session_state[f"{template_name}_groups_editable"][next_index], \
                                st.session_state[f"{template_name}_groups_editable"][group_index]
                            st.rerun()

            with content_col:
                if is_editable:
                    # Create a header container with flexbox layout
                    st.markdown(
                        f"""
                        <div class="group-header">
                            <p class="smaller-font group-header-text">Editable Group {sum(1 for x in st.session_state[f'{template_name}_groups_editable'][:group_index] if x) + 1}:</p>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

                    # If group is empty, show add line button
                    if not current_group:
                        if st.button("+ Add New Line", key=f"add_line_{group_index}"):
                            st.session_state[f"{template_name}_groups"][group_index] = [
                                "New line"]
                            st.rerun()
                    else:
                        # Show grouped text area for editable content
                        group_text = "\n".join(current_group)
                        # Calculate height based on content, with minimum height of 68 pixels
                        num_lines = len(current_group)
                        # Ensure minimum height of 68px
                        height = max(68, num_lines * 25 + 25)

                        col1, col2 = st.columns([0.95, 0.05])
                        with col1:
                            st.text_area(
                                label=f"Edit group {group_index + 1}",
                                value=group_text,
                                key=f"{template_name}_group_{group_index}",
                                height=height,
                                label_visibility="collapsed",
                                on_change=on_text_change,
                                args=(template_name, group_index,)
                            )

                        with col2:
                            # Add delete button for the group
                            if st.button("🗑", key=f"delete_group_{group_index}", help="Delete entire group", type="secondary"):
                                st.session_state[f"{template_name}_groups"][group_index] = [
                                ]
                                st.rerun()
                else:
                    # Handle non-editable line as a parent element
                    if current_group:  # Only process if group is not empty
                        content_col, delete_col = st.columns([0.95, 0.05])

                        with content_col:
                            # Show the line with any dynamic content highlighted
                            # Since each non-editable line is now its own group
                            line = current_group[0]
                            dynamic_found = any(
                                dyn in line for dyn in dynamic_values)

                            if dynamic_found:
                                st.code(line, language="python")
                            else:
                                st.text(line)

                        with delete_col:
                            # Delete button for the line
                            if st.button("🗑", key=f"delete_line_{group_index}", help="Delete line"):
                                groups_to_remove.add(group_index)
                                st.session_state[f"{template_name}_deleted_groups"].add(
                                    group_index)
                                st.rerun()

            # Add lines from current group if not marked for removal
            if group_index not in groups_to_remove and group_index not in st.session_state[f"{template_name}_deleted_groups"]:
                edited_lines.extend(current_group)
            st.markdown("---")

        # Update the complete template
        if edited_lines:
            st.session_state.edited_templates[template_name] = "\n".join(
                edited_lines).rstrip()

    return edited_lines


def load_json_file(file_path):
    try:
        with open(file_path, 'r') as f:
            return json.load(f)
    except Exception as e:
        st.error(f"Error loading JSON file: {str(e)}")
        return None


def get_json_files():
    json_dir = "simplifyJson"
    json_files = []
    try:
        for file in os.listdir(json_dir):
            if file.endswith(".json"):
                json_files.append(file)
        return sorted(json_files)
    except Exception as e:
        st.error(f"Error accessing simplifyJson directory: {str(e)}")
        return []


def read_config_from_s3(config_name):
    """Read a config file from S3"""
    try:
        # Get the object from S3
        response = s3_client.get_object(
            Bucket=S3_BUCKET,
            Key=f"clause_configs/{config_name}.py"
        )
        content = response['Body'].read().decode('utf-8')

        # Create a temporary module namespace
        namespace = {}
        exec(content, namespace)

        # Find the first uppercase variable which should be our config dictionary
        config_dict = next((val for name, val in namespace.items()
                            if name.isupper() and isinstance(val, dict)), {})

        return config_dict
    except Exception as e:
        st.error(f"Error reading config from S3: {str(e)}")
        return {}


def get_config_files_from_s3():
    """List all config files in S3"""
    try:
        response = s3_client.list_objects_v2(
            Bucket=S3_BUCKET,
            Prefix='clause_configs/'
        )

        config_files = []
        for obj in response.get('Contents', []):
            filename = os.path.basename(obj['Key'])
            if filename.endswith('_config.py') and not filename.startswith('__'):
                config_files.append(filename[:-3])  # Remove .py extension
        return sorted(config_files)
    except Exception as e:
        st.error(f"Error listing configs from S3: {str(e)}")
        return []


def split_prompt_template(template):
    # Split by newlines while preserving them
    lines = template.split('\n')
    # Find dynamic content (anything between {})
    dynamic_parts = re.finditer(r'{[^}]+}', template)
    dynamic_values = [match.group() for match in dynamic_parts]
    return lines, dynamic_values


def get_config_files():
    config_dir = "clause_configs"
    config_files = []
    for file in os.listdir(config_dir):
        if file.endswith("_config.json") and not file.startswith("__"):
            config_files.append(file[:-5])  # Remove .json extension
    return sorted(config_files)


def save_config_changes_to_s3(config_name, edited_templates):
    """Save changes back to S3 using Python format"""
    try:
        logger.info(
            f"Debug: Attempting to save changes to S3 for {config_name}")
        logger.info(f"Debug: Edited templates: {edited_templates}")
        logger.info(f"Debug: Using S3 bucket: {S3_BUCKET}")
        logger.info(
            f"Debug: AWS credentials present: {bool(os.getenv('AWS_ACCESS_KEY_ID') and os.getenv('AWS_SECRET_ACCESS_KEY'))}")

        # First read the existing file from S3
        response = s3_client.get_object(
            Bucket=S3_BUCKET,
            Key=f"clause_configs/{config_name}.py"
        )
        content = response['Body'].read().decode('utf-8')
        logger.info("Debug: Successfully read existing file from S3")

        # Create a temporary module namespace
        namespace = {}
        exec(content, namespace)
        logger.info("Debug: Successfully executed existing content")

        # Find the config dictionary and its name
        config_var_name = next((name for name in namespace if name.isupper(
        ) and isinstance(namespace[name], dict)), None)
        if not config_var_name:
            st.error("Could not find config dictionary in module")
            logger.error("Could not find config dictionary in module")
            return False

        logger.info(f"Debug: Found config variable name: {config_var_name}")
        config_dict = namespace[config_var_name]

        # Update only the prompt_template in the dictionary while preserving all other fields
        for key, new_template in edited_templates.items():
            if key in config_dict:
                config_dict[key]['prompt_template'] = new_template
                logger.info(f"Debug: Updated template for {key}")
            else:
                error_msg = f"Warning: Key {key} not found in config"
                st.error(error_msg)
                logger.error(error_msg)
                return False

        # Create the new file content
        file_content = [
            "# Auto-generated config file",
            f"{config_var_name} = {{"
        ]

        # Process each template
        for i, (key, value) in enumerate(config_dict.items()):
            # Start template entry
            file_content.append(f'    "{key}": {{')

            # Process each field in the template
            field_lines = []
            for field_key, field_value in value.items():
                if field_key == 'prompt_template':
                    # Handle multi-line prompt template
                    field_lines.append(
                        f'        "prompt_template": """\n{field_value}\n"""')
                elif isinstance(field_value, str):
                    field_lines.append(
                        f'        "{field_key}": "{field_value}"')
                elif isinstance(field_value, bool):
                    # Keep as Python bool (True/False)
                    field_lines.append(
                        f'        "{field_key}": {field_value}')
                elif isinstance(field_value, (int, float)):
                    field_lines.append(f'        "{field_key}": {field_value}')
                elif field_value is None:
                    field_lines.append(f'        "{field_key}": None')
                elif isinstance(field_value, (list, dict)):
                    # Use repr for lists and dicts to maintain Python syntax
                    field_lines.append(
                        f'        "{field_key}": {repr(field_value)}')
                else:
                    field_lines.append(
                        f'        "{field_key}": {repr(field_value)}')

            # Join fields with commas
            file_content.append(',\n'.join(field_lines))

            # Close template entry
            if i < len(config_dict) - 1:
                file_content.append('    },')
            else:
                file_content.append('    }')

        # Close main dictionary
        file_content.append("}")

        # Join all lines with proper newlines
        final_content = '\n'.join(file_content)
        logger.info("Debug: Generated new file content")

        # Upload the modified content back to S3
        logger.info(f"Debug: Attempting to upload to S3")
        try:
            s3_client.put_object(
                Bucket=S3_BUCKET,
                Key=f"clause_configs/{config_name}.py",
                Body=final_content.encode('utf-8')
            )
            logger.info(f"Debug: Save completed successfully")
            return True
        except Exception as s3_error:
            logger.error(f"Error uploading to S3: {str(s3_error)}")
            raise

    except Exception as e:
        error_msg = f"Error saving changes to S3: {str(e)}"
        st.error(error_msg)
        logger.error(error_msg)
        return False


def run_summary_generation(json_data, config_dict, selected_template=None):
    try:
        summary_outputs = []
        prompt_logs = []  # Store prompts for logging

        # If selected_template is provided, only process that template
        if selected_template:
            templates_to_process = {
                selected_template: config_dict[selected_template]}
        else:
            templates_to_process = config_dict

        for clause_name, clause_config in templates_to_process.items():
            # If we have edited this template, use the edited version
            if clause_name in st.session_state.edited_templates:
                clause_config = clause_config.copy()
                clause_config['prompt_template'] = st.session_state.edited_templates[clause_name]

            result = process_clause_config(clause_config, json_data)
            if result["output"] and result["output"] != "No output generated.":
                if (result.get("summary_type", "").lower() == "concise" and
                        clause_config.get("view_prompt", True) is False):
                    continue

                # Store the prompt and result for logging
                prompt_logs.append({
                    "clause_name": clause_name,
                    "template": clause_config['prompt_template'],
                    "final_prompt": result.get("used_prompt", "Prompt not available"),
                    "output": result["output"]
                })

                summary_outputs.append({
                    "clause_name": clause_name,
                    **result
                })

        # Sort summaries by rank before writing
        summary_outputs_sorted = sorted(
            summary_outputs, key=lambda x: x.get('summary_rank', 999))

        # Generate unique output path for this run
        output_path = f"clause_summary_{selected_template}.docx" if selected_template else "clause_summary_output.docx"

        try:
            # Create a new document
            doc = Document()

            # Add a simple title
            doc.add_heading('Merger Agreement Clause Summaries', 0)

            # Add each summary
            for summary in summary_outputs_sorted:
                # Add section heading if needed
                if summary.get("summary_display_section"):
                    doc.add_heading(
                        summary["summary_display_section"], level=1)

                # Add the summary text
                p = doc.add_paragraph()
                p.add_run("• ").bold = True
                p.add_run(summary["output"])

                # Add references if any
                if summary.get("references"):
                    ref_para = doc.add_paragraph()
                    ref_para.add_run("References: " +
                                     "; ".join(summary["references"]))
                    ref_para.style = 'List Bullet'

            # Save the document
            doc.save(output_path)

            if os.path.exists(output_path):
                return True, prompt_logs, output_path
            else:
                st.error("Failed to write the document")
                return False, prompt_logs, None

        except Exception as e:
            st.error(f"Error writing document: {str(e)}")
            return False, prompt_logs, None

    except Exception as e:
        st.error(f"Error generating summary: {str(e)}")
        return False, [], None


def group_lines_by_editability(lines, dynamic_values):
    """Group lines into editable and non-editable sections"""
    groups = []
    current_group = {"editable": True, "lines": [], "indices": []}

    # First, clean up the lines by removing trailing empty lines
    while lines and not lines[-1].strip():
        lines.pop()

    for i, line in enumerate(lines):
        has_dynamic = any(dyn in line for dyn in dynamic_values)
        is_editable = not has_dynamic

        if is_editable:
            # If current group is editable, add to it
            if current_group["editable"]:
                current_group["lines"].append(line)
                current_group["indices"].append(i)
            else:
                # If current group is not editable, save it and start new editable group
                if current_group["lines"]:
                    groups.append(current_group)
                current_group = {"editable": True,
                                 "lines": [line], "indices": [i]}
        else:
            # For non-editable lines, create a separate group for each line
            if current_group["lines"]:
                groups.append(current_group)
            # Create a new group for this single non-editable line
            groups.append({"editable": False, "lines": [line], "indices": [i]})
            # Start a new editable group
            current_group = {"editable": True, "lines": [], "indices": []}

    # Add the last group only if it has non-empty lines
    if current_group["lines"] and any(line.strip() for line in current_group["lines"]):
        groups.append(current_group)

    return groups


def on_save_click():
    """Callback for save button click"""
    logger.info("Save button clicked")
    try:
        logger.info("Attempting save operation")

        # Store save data
        save_data = {
            'config': selected_config,
            'template': st.session_state.selected_template,
            'content': st.session_state.edited_templates[st.session_state.selected_template]
        }
        logger.info(f"Using save data: {save_data}")

        # Attempt to save
        success = save_config_changes_to_s3(
            save_data['config'],
            {save_data['template']: save_data['content']}
        )

        if success:
            logger.info("Save completed successfully")
            # Remove from edited templates
            del st.session_state.edited_templates[save_data['template']]
            logger.info("Session state cleared after successful save")
            # Set flag for rerun instead of calling rerun directly
            st.session_state.needs_rerun = True
        else:
            logger.error("Save operation failed")
            st.error("❌ Failed to save changes to S3")
    except Exception as e:
        logger.error(f"Error during save operation: {str(e)}")
        st.error(f"❌ Error saving to S3: {str(e)}")
        import traceback
        logger.error(f"Full error traceback: {traceback.format_exc()}")
        st.code(traceback.format_exc())


# Custom CSS for smaller headers
st.markdown("""
    <style>
    .small-font {
        font-size: 18px !important;
        margin-bottom: 0px !important;
    }
    .medium-font {
        font-size: 20px !important;
        margin-bottom: 0px !important;
        font-weight: bold;
    }
    .large-font {
        font-size: 28px !important;
        margin-bottom: 0px !important;
        font-weight: bold;
    }
            
    .pt-20{
        padding-top: 20px !important;
    }
    .pb-10{
        padding-bottom: 10px !important;
    }
    .smaller-font {
        font-size: 16px !important;
        margin-bottom: 12px !important;
        font-weight: bold;
    }
    .smallest-font {
        font-size: 14px !important;
        margin-bottom: 0px !important;
    }
    /* Reduce padding in Streamlit elements */
    .st-emotion-cache-z5fcl4 {
        padding-top: 1rem !important;
        padding-bottom: 1rem !important;
    }
   
    .st-emotion-cache-16idsys {
        padding-top: 0.5rem !important;
        padding-bottom: 0.5rem !important;
    }
    .st-emotion-cache-16txtl3 {
        padding-top: 0.5rem !important;
        padding-bottom: 0.5rem !important;
    }
    .st-emotion-cache-16idsys p {
        margin-bottom: 0.5rem !important;
    }
    /* Reduce padding in expander */
    .streamlit-expanderHeader {
        padding-top: 0.5rem !important;
        padding-bottom: 0.5rem !important;
    }
    /* Reduce padding in text areas */
    .stTextArea {
        padding-top: 0.1rem !important;
        padding-bottom: 0.1rem !important;
    }
    /* Reduce margins between elements */
    div.stButton > button {
        margin-top: 0.5rem !important;
        margin-bottom: 0.5rem !important;
    }
    .block-container {
        padding-top: 2rem !important;
        padding-bottom: 2rem !important;
    }
             
    </style>
""", unsafe_allow_html=True)

# st.markdown('<p class="large-font">Prompt Template Editor</p>',
#             unsafe_allow_html=True)

# Move steps 1-3 to sidebar
with st.sidebar:
    st.markdown('<p class="medium-font">Prompt Template Editor</p>',
                unsafe_allow_html=True)
    st.markdown("---")  # Add a divider line

    # JSON File Selection (Step 1)
    st.markdown('<p class="smaller-font pt-20"><u>1. Select JSON Source</u></p>',
                unsafe_allow_html=True)
    json_source = st.radio("Choose JSON source:", [
                           "Upload JSON", "Select Existing JSON"])

    json_data = None
    if json_source == "Upload JSON":
        uploaded_file = st.file_uploader("Upload JSON file", type=['json'])
        if uploaded_file:
            json_data = json.load(uploaded_file)
            st.success("✅ JSON file loaded successfully!")
    else:
        # Show dropdown for existing JSON files
        json_files = get_json_files()
        if json_files:
            selected_json = st.selectbox("Choose a JSON file:", json_files)
            if selected_json:
                json_path = os.path.join("simplifyJson", selected_json)
                json_data = load_json_file(json_path)
                if json_data:
                    st.success("✅ JSON file loaded successfully!")
        else:
            st.warning("No JSON files found in simplifyJson directory")

    # Config File Selection (Step 2)
    st.markdown('<p class="smaller-font"><u>2. Select Config File</u></p>',
                unsafe_allow_html=True)
    config_files = get_config_files_from_s3()
    selected_config = st.selectbox("Choose a config file:", config_files)

    # Template Selection (Step 3)
    if selected_config and json_data:
        config_dict = read_config_from_s3(selected_config)
        st.markdown('<p class="smaller-font"><u>3. Select Template to Edit</u></p>',
                    unsafe_allow_html=True)
        template_options = list(config_dict.keys())
        st.session_state.selected_template = st.selectbox(
            "Choose a template to edit:", template_options)

# Full width section for Step 4 (Editing) in main content area
if st.session_state.selected_template:
    st.markdown('<p class="smaller-font"><u>4. Edit Selected Template</u></p>',
                unsafe_allow_html=True)

    value = config_dict[st.session_state.selected_template]
    if isinstance(value, dict) and "prompt_template" in value:
        prompt_template = value["prompt_template"]

        # Split the template into lines and find dynamic content
        lines, dynamic_values = split_prompt_template(prompt_template)

        # Group lines by editability
        groups = group_lines_by_editability(lines, dynamic_values)

        # Use the new render_editable_groups function
        edited_lines = render_editable_groups(
            st.session_state.selected_template, groups, dynamic_values)

        # Show the complete edited template
        st.markdown(
            '<p class="smallest-font pb-10">Complete Template:</p>', unsafe_allow_html=True)
        # Convert whitespace to visible characters and wrap in <pre> tag for preserving formatting
        template_with_whitespace = "\n".join(
            edited_lines).rstrip()  # Remove trailing whitespace

        st.code(template_with_whitespace, language="plaintext")

        # Update the complete template without trailing newlines
        if edited_lines:
            st.session_state.edited_templates[st.session_state.selected_template] = "\n".join(
                edited_lines).rstrip()

        # Add run button
        if st.button("Run Summary"):
            if json_data and selected_config:
                with st.spinner(f"Generating summary for {st.session_state.selected_template}..."):
                    success, prompt_logs, output_path = run_summary_generation(
                        json_data, config_dict, st.session_state.selected_template)

                    if success and output_path and os.path.exists(output_path):
                        st.success(f"✅ Summary generated successfully!")

                        # Create tabs for Results and Changes
                        results_tab, changes_tab = st.tabs(
                            ["Summary Results", "Prompt Changes"])

                        with results_tab:
                            try:
                                # Create a container for the download link
                                download_container = st.container()

                                with download_container:
                                    if os.path.exists(output_path):
                                        # Read file and encode to base64
                                        with open(output_path, "rb") as file:
                                            file_bytes = file.read()
                                            import base64
                                            b64_data = base64.b64encode(
                                                file_bytes).decode()

                                            file_name = os.path.basename(
                                                output_path)
                                            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

                                            # Create HTML download link
                                            href = f'<a href="data:{mime_type};base64,{b64_data}" download="{file_name}" style="text-decoration: none;"><div style="background-color: #a3a3a2; color: white; padding: 8px 16px; border-radius: 4px; cursor: pointer; display: inline-block; text-align: center;">📥 Download Summary</div></a>'
                                            st.markdown(
                                                href, unsafe_allow_html=True)

                                # Show prompts used in generation
                                if prompt_logs:
                                    st.markdown(
                                        '<p class="smaller-font">📋 Prompt Used in Generation</p>', unsafe_allow_html=True)
                                    log = prompt_logs[0]  # Show first log
                                    with st.expander(f"**{log['clause_name']}**"):
                                        st.markdown(
                                            '<p class="smallest-font">Final Prompt</p>', unsafe_allow_html=True)
                                        st.code(log['final_prompt'])

                                        st.markdown(
                                            '<p class="smallest-font">Generated Output</p>', unsafe_allow_html=True)
                                        st.code(log['output'])

                            except Exception as e:
                                st.error(
                                    f"Error preparing download: {str(e)}")

                        with changes_tab:
                            logger.info("Entering changes_tab section")

                            # Check if we need to rerun from previous save
                            if st.session_state.needs_rerun:
                                st.session_state.needs_rerun = False
                                st.rerun()

                            if st.session_state.selected_template in st.session_state.edited_templates:
                                logger.info(
                                    f"Found template {st.session_state.selected_template} in edited_templates")
                                original = config_dict[st.session_state.selected_template]["prompt_template"]
                                edited = st.session_state.edited_templates[
                                    st.session_state.selected_template]

                                st.markdown("### Template Comparison")

                                st.markdown("**Original Template:**")
                                st.text_area(
                                    "Original", value=original, height=300, disabled=True, label_visibility="collapsed")

                                st.markdown("**Modified Template:**")
                                st.text_area(
                                    "Modified", value=edited, height=300, disabled=True, label_visibility="collapsed")

                                # Create save button
                                st.button(
                                    "💾 Save Changes", key="save_button", on_click=on_save_click)
                            else:
                                logger.info(
                                    f"Template {st.session_state.selected_template} not found in edited_templates")
    else:
        st.warning(
            "No prompt template found in this configuration.")
elif not json_data:
    st.info("Please select a JSON file first")
elif not selected_config:
    st.info("Please select a config file")
